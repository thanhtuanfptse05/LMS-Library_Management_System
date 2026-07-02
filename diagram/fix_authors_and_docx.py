import re
import sys
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, val in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(val))

def build_docx():
    # 1. Parse spec file to map UCs to owners
    uc_owner = {}
    with open('spec-UC-BR-FR.txt', 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    current_owner = "Unknown"
    for line in lines:
        if line.startswith('//'): continue
        if '- Owner:' in line:
            current_owner = line.split(':')[1].strip()
        elif '- UC:' in line:
            ucs = re.findall(r'UC-\d+', line)
            for uc in ucs:
                uc_owner[uc] = current_owner

    # Fallback missing ones if any
    
    # 2. Update detailed-UC-specifications.md
    with open('detailed-UC-specifications.md', 'r', encoding='utf-8') as f:
        content = f.read()
        
    tables = re.findall(r'<table[^>]*>(.*?)</table>', content, re.DOTALL)
    
    new_content = content
    for table_str in tables:
        # Find UC ID
        uc_match = re.search(r'<b>(UC-\d+):.*?</b>', table_str)
        if uc_match:
            uc_id = uc_match.group(1)
            owner = uc_owner.get(uc_id, "Tuan")
            
            # Replace author in table
            # Original looks like: <td><b>Created By:</b></td>\n    <td width="30%">Tuan</td>
            # Use regex to replace the cell next to Created By:
            row1_pattern = r'(<td><b>Created By:</b></td>\s*<td[^>]*>)(.*?)(</td>)'
            
            def repl_author(m):
                return m.group(1) + owner + m.group(3)
                
            new_table_str = re.sub(row1_pattern, repl_author, table_str, count=1)
            if new_table_str != table_str:
                new_content = new_content.replace(table_str, new_table_str)
                
    with open('detailed-UC-specifications.md', 'w', encoding='utf-8') as f:
        f.write(new_content)

    # 3. Create DOCX
    doc = Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    doc.add_heading('Library Management System (LMS) - Use Case Specifications', 0)
    
    tables_again = re.findall(r'<table[^>]*>(.*?)</table>', new_content, re.DOTALL)
    
    uc_count = 1
    
    for table_idx, table_str in enumerate(tables_again):
        rows_data = []
        row_pattern = re.compile(r'<tr>(.*?)</tr>', re.DOTALL)
        rows = row_pattern.findall(table_str)
        
        # Determine the number of rows in this table
        num_rows = len(rows)
        
        # Add heading above table
        # Extract UC name from first row
        first_row_cells = re.findall(r'<td[^>]*>(.*?)</td>', rows[0], re.DOTALL)
        if len(first_row_cells) >= 2:
            uc_title = re.sub(r'<[^>]*>', '', first_row_cells[1]).strip()
            # Handle special characters and &rarr; if any, though unlikely in title
            uc_title = uc_title.replace('&rarr;', '->').replace('&lt;', '<').replace('&gt;', '>')
            
            heading = doc.add_heading(f"1.{uc_count} {uc_title}", level=2)
            # You can style the heading if needed
        uc_count += 1
        
        docx_table = doc.add_table(rows=num_rows, cols=4)
        docx_table.autofit = False
        
        widths = [Inches(1.5), Inches(2.0), Inches(1.5), Inches(2.0)]
        for i, col in enumerate(docx_table.columns):
            col.width = widths[i]
            
        for row_idx, row_content in enumerate(rows):
            cell_pattern = re.compile(r'<td[^>]*>(.*?)</td>', re.DOTALL)
            cells = cell_pattern.findall(row_content)
            
            def clean_text(text):
                text = re.sub(r'<b>(.*?)</b>', r'\1', text)
                text = text.replace('<br>', '\n').replace('<br/>', '\n').replace('&rarr;', '->').replace('&lt;', '<').replace('&gt;', '>')
                text = re.sub(r'<[^>]*>', '', text)
                return text.strip()
                
            if row_idx == 0:
                h = clean_text(cells[0])
                v = clean_text(cells[1])
                docx_table.cell(0, 0).text = h
                merged = docx_table.cell(0, 1).merge(docx_table.cell(0, 3))
                merged.text = v
            elif row_idx in (1, 2):
                h1 = clean_text(cells[0])
                v1 = clean_text(cells[1])
                h2 = clean_text(cells[2])
                v2 = clean_text(cells[3])
                docx_table.cell(row_idx, 0).text = h1
                docx_table.cell(row_idx, 1).text = v1
                docx_table.cell(row_idx, 2).text = h2
                docx_table.cell(row_idx, 3).text = v2
            else:
                h = clean_text(cells[0])
                v = clean_text(cells[1])
                docx_table.cell(row_idx, 0).text = h
                merged = docx_table.cell(row_idx, 1).merge(docx_table.cell(row_idx, 3))
                merged.text = v
                
        for r_idx in range(num_rows):
            for c_idx in range(4):
                cell = docx_table.cell(r_idx, c_idx)
                set_cell_border(
                    cell,
                    top={"sz": 12, "val": "single", "color": "000000", "space": "0"},
                    bottom={"sz": 12, "val": "single", "color": "000000", "space": "0"},
                    left={"sz": 12, "val": "single", "color": "000000", "space": "0"},
                    right={"sz": 12, "val": "single", "color": "000000", "space": "0"}
                )
                
                # Make header cells bold
                is_header = (c_idx == 0) or (r_idx in (1, 2) and c_idx == 2)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(10.5)
                        if is_header:
                            run.font.bold = True
                            
        doc.add_paragraph('\n')
        
    doc.save('detailed-UC-specifications.docx')
    print("Success! Created detailed-UC-specifications.docx")

if __name__ == '__main__':
    build_docx()
