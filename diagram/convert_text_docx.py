import re
from docx import Document

def build_docx():
    doc = Document()
    
    with open('detailed-UC-specifications.md', 'r', encoding='utf-8') as f:
        content = f.read()
        
    table_pattern = re.compile(r'<table[^>]*>(.*?)</table>', re.DOTALL)
    tables = table_pattern.findall(content)
    
    uc_count = 1
    
    for table_idx, table_str in enumerate(tables):
        rows_data = []
        row_pattern = re.compile(r'<tr>(.*?)</tr>', re.DOTALL)
        rows = row_pattern.findall(table_str)
        
        cells_data = []
        
        for row_idx, row_content in enumerate(rows):
            cell_pattern = re.compile(r'<td[^>]*>(.*?)</td>', re.DOTALL)
            cells = cell_pattern.findall(row_content)
            
            def clean_text(text):
                text = re.sub(r'<b>(.*?)</b>', r'\1', text)
                text = text.replace('<br>', '\n').replace('<br/>', '\n').replace('&rarr;', '->').replace('&lt;', '<').replace('&gt;', '>')
                text = re.sub(r'<[^>]*>', '', text)
                return text.strip()
                
            if row_idx == 0:
                h = clean_text(cells[0])
                v = clean_text(cells[1])
                cells_data.append((h, v))
                
                # Heading
                doc.add_heading(f"1.{uc_count} {v}", level=2)
                uc_count += 1
            elif row_idx in (1, 2):
                h1 = clean_text(cells[0])
                v1 = clean_text(cells[1])
                h2 = clean_text(cells[2])
                v2 = clean_text(cells[3])
                cells_data.append((h1, v1))
                cells_data.append((h2, v2))
            else:
                h = clean_text(cells[0])
                v = clean_text(cells[1])
                cells_data.append((h, v))
                
        for h, v in cells_data:
            doc.add_paragraph(h)
            doc.add_paragraph(v)
            
    doc.save('detailed-UC-specifications.docx')
    print("Success! Created detailed-UC-specifications.docx with plain text layout")

if __name__ == '__main__':
    build_docx()
