import re
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, val in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(val))

def build_docx():
    doc = Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    doc.add_heading('Library Management System (LMS) - Use Case Specifications', 0)
    
    with open('detailed-UC-specifications.md', 'r', encoding='utf-8') as f:
        content = f.read()
        
    table_pattern = re.compile(r'<table[^>]*>(.*?)</table>', re.DOTALL)
    tables = table_pattern.findall(content)
    
    for table_idx, table_str in enumerate(tables):
        rows_data = []
        row_pattern = re.compile(r'<tr>(.*?)</tr>', re.DOTALL)
        rows = row_pattern.findall(table_str)
        
        docx_table = doc.add_table(rows=14, cols=4)
        docx_table.autofit = False
        
        widths = [Inches(1.5), Inches(2.0), Inches(1.5), Inches(2.0)]
        for i, col in enumerate(docx_table.columns):
            col.width = widths[i]
            
        for row_idx, row_content in enumerate(rows):
            cell_pattern = re.compile(r'<td[^>]*>(.*?)</td>', re.DOTALL)
            cells = cell_pattern.findall(row_content)
            
            def clean_text(text):
                text = re.sub(r'<b>(.*?)</b>', r'\1', text)
                text = text.replace('<br>', '\n').replace('<br/>', '\n').replace('&rarr;', '->').replace('&lt;', '<').replace('&gt;', '>')
                text = re.sub(r'<[^>]*>', '', text)
                return text.strip()
                
            if row_idx == 0:
                h = clean_text(cells[0])
                v = clean_text(cells[1])
                docx_table.cell(0, 0).text = h
                merged = docx_table.cell(0, 1).merge(docx_table.cell(0, 3))
                merged.text = v
            elif row_idx in (1, 2):
                h1 = clean_text(cells[0])
                v1 = clean_text(cells[1])
                h2 = clean_text(cells[2])
                v2 = clean_text(cells[3])
                docx_table.cell(row_idx, 0).text = h1
                docx_table.cell(row_idx, 1).text = v1
                docx_table.cell(row_idx, 2).text = h2
                docx_table.cell(row_idx, 3).text = v2
            else:
                h = clean_text(cells[0])
                v = clean_text(cells[1])
                docx_table.cell(row_idx, 0).text = h
                merged = docx_table.cell(row_idx, 1).merge(docx_table.cell(row_idx, 3))
                merged.text = v
                
        for r_idx in range(14):
            for c_idx in range(4):
                cell = docx_table.cell(r_idx, c_idx)
                set_cell_border(
                    cell,
                    top={"sz": 12, "val": "single", "color": "000000", "space": "0"},
                    bottom={"sz": 12, "val": "single", "color": "000000", "space": "0"},
                    left={"sz": 12, "val": "single", "color": "000000", "space": "0"},
                    right={"sz": 12, "val": "single", "color": "000000", "space": "0"}
                )
                
                # Make header cells bold
                is_header = (c_idx == 0) or (r_idx in (1, 2) and c_idx == 2)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(10.5)
                        if is_header:
                            run.font.bold = True
                            
        doc.add_paragraph('\n')
        
    doc.save('detailed-UC-specifications.docx')
    print("Success! Created detailed-UC-specifications.docx")

if __name__ == '__main__':
    build_docx()
